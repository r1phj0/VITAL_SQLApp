import xlwings as xw
import pandas as pd
import dateparser
from connect_to_db import get_connection
from tkinter import simpledialog
from GeneralMethods import ScheduleQuery, StopQuery, ScheduleValidation, stop_time_mismatch, query_rot
from SQLDataAppMethods import get_times_schedule_df
import time
import numpy as np
from docx import Document
from docx.shared import RGBColor


def create_schedule_objects(post_zb_all, dates):
    schedule_objects = []
    x = 0
    while x < len(post_zb_all['Schedule #']):
        try:
            schedule_objects.append(ScheduleValidation(post_zb_all.loc[x, :], dates))
        except:
            print(x)
            print('skip it?')
        x += 1
    schedule_dict = {k.schedule_num: k for k in schedule_objects}
    return schedule_dict


def determine_date(type, dates):
    date = ''
    if type == 'D':
        date = dates[0]
        if date == '':
            date = simpledialog.askstring('D Date', 'Please input the correct D date:')
    elif type == 'DS':
        type = 'D'
        date = dates[1]
        if date == '':
            date = simpledialog.askstring('DS Date', 'Please input the correct DS date:')
    elif type == 'F':
        date = dates[2]
        if date == '':
            date = simpledialog.askstring('F Date', 'Please input the correct F date:')
    return date, type


# pulls schedules all at once by type, not schedule number
def print_sched_dfs(types, dates, conn, df_all, site):
    all_types_dataframes = []
    for type in types:
        # isolates dataframe where the schedule type is equal to the type from the above line
        df_type = df_all[(df_all['Type (D / DS / I / F)'] == type)]
        # isolates schedule numbers with that type
        schedule_nums = df_type['Schedule #'].tolist()
        # gets the date from the current type
        date, type = determine_date(type, dates)
        # creating schedule query object
        new_SchedQuery = ScheduleQuery()
        final_query, p = new_SchedQuery.build_zb_spec_query(site, type, date, schedule_nums)
        final_db = pd.read_sql(final_query, conn, params=p)
        all_types_dataframes.append(final_db)
    final_dataframe = pd.concat(all_types_dataframes)
    if final_dataframe.empty:
        print("Schedules Dataframe is empty")
    final_dataframe.to_excel('DataFiles\\' + site + ' SCHEDULES ZB ' + time.strftime("%Y%m%d") + '.xlsx', index=False)
    final_dataframe = final_dataframe.reset_index(drop=True)
    return final_dataframe


def print_stop_dfs(conn, sched_df, type_dict, site, schedule_obj_dict):
    stops_not_found = ''
    all_stops_dataframes = []
    for sched in sched_df['SCH_SCHED_NBR'].tolist():
        new_StopQuery = StopQuery()
        final_stop_query = new_StopQuery.build_zb_spec_stop_query(site, sched, schedule_obj_dict[sched].zb_type,
                                                                  schedule_obj_dict[sched].date)
        final_stop_db = pd.read_sql(final_stop_query, conn)
        if final_stop_db.empty:
            stops_not_found = stops_not_found + sched + ', '
        else:
            all_stops_dataframes.append(final_stop_db)
    if stops_not_found != '':
        print('No matching stops for the following schedules: ', stops_not_found)

    final_dataframe = pd.concat(all_stops_dataframes).drop_duplicates()
    final_dataframe.to_excel('DataFiles\\' + site + ' STOPS ZB ' + time.strftime("%Y%m%d") + '.xlsx', index=False)
    return final_dataframe


def sort_and_print_sched_stops(schedule_numbers_file):
    # loading in workbook mileage & hours sheet
    workbook = xw.Book(schedule_numbers_file)
    sheet = workbook.sheets['Miles & Hours']

    # site name will always be in cell A1
    site = sheet['A1:A1'].options(pd.DataFrame, index=False, header=False).value
    site = site[0].tolist()[0]

    # getting everything from Miles & Hours post ZB
    post_zb_all = sheet['J15:R1000'].options(pd.DataFrame, index=False, header=True).value
    post_zb_all = post_zb_all.dropna()
    post_zb_all = post_zb_all.reset_index()
    if 'index' in post_zb_all.columns.tolist():
        post_zb_all = post_zb_all.drop('index', axis=1)

    post_zb_all['Schedule #'] = post_zb_all['Schedule #'].map(str)
    post_zb_all['Schedule #'] = post_zb_all['Schedule #'].apply(lambda x: x.replace('.0',''))

    post_zb_w_hours = post_zb_all[post_zb_all['Paid Dly Hrs']!= "00:00"].reset_index(drop=True)

    # dropping columns
    sched_w_type = post_zb_w_hours[['Type (D / DS / I / F)', 'Schedule #']]

    # getting dates for drafts/future/etc
    df_dates = sheet['U6:U8'].options(pd.DataFrame, index=False, header=False).value
    df_dates = df_dates.fillna("01/01/1111")
    # zero represents first column, so putting all possible dates into a list
    dates = df_dates[0].tolist()

    # setting dates as correct format for query
    if len(dates) != 0:
        dates = [dateparser.parse(x).strftime('%d-%m-%Y') for x in dates if x != '']
    types = list(set(sched_w_type['Type (D / DS / I / F)']))
    types_dates = {}
    if 'I' in types:
        types_dates['I'] = ""
    if 'D' in types:
        types_dates['D'] = dates[0]
    if 'DS' in types:
        types_dates['DS'] = dates[1]
    if 'F' in types:
        types_dates['F'] = dates[2]
    workbook.app.quit()

    # creating dict with schedule number as the key and a ScheduleValidation object from General methods as the value
    schedule_obj_dict = create_schedule_objects(post_zb_all, dates)

    # establishing connection before pulling schedules and stops files
    conn = get_connection()
    print('Getting schedules...')
    type_dict = {str(sched_w_type['Schedule #'].tolist()[x]): sched_w_type['Type (D / DS / I / F)'].tolist()[x] for x in
                 range(len(sched_w_type['Type (D / DS / I / F)'].tolist()))}

    schedule_df = print_sched_dfs(types, dates, conn, sched_w_type, site)

    # getting times route for ROT df
    data_times_for_rot = get_times_schedule_df(schedule_df, conn)
    data_times_for_rot.to_excel('DataFiles\\' + site + ' TIMES ROUTE ZB ' + time.strftime("%Y%m%d") + '.xlsx', index=False)

    print('Getting stops...')
    stops_df = print_stop_dfs(conn, schedule_df, type_dict, site, schedule_obj_dict)

    print("Getting ROT")
    p = {"site0": site}
    rot_query = query_rot.replace('SITES_STRING_TO_REPLACE', ":site0")
    data_rot = pd.read_sql(rot_query, conn, params =p)
    if data_rot.empty:
        print("Oh no! ROT file for ", site, " is empty!")
    data_rot.to_excel('DataFiles\\' + site + ' ROT ZB ' + time.strftime("%Y%m%d") + '.xlsx', index=False)

    # getting all schedule numbers that were printed
    printed_schedules = [str(sch_nbr) for sch_nbr in schedule_df['SCH_SCHED_NBR']]
    # getting all schedule numbers that were in the miles & hours tab that are not in the printed schedules from sql
    not_found_schedules = [str(x) for x in sched_w_type['Schedule #'].tolist() if str(x) not in printed_schedules]
    missing_schedule_data = pd.DataFrame()
    if len(not_found_schedules) > 0:
        missing_schedule_data = process_missing_schedules(site, not_found_schedules, conn)
    print('Schedules not found: ', not_found_schedules)

    # checking if there are multiple of 1 schedule number (have to have the parameter in case different eff. dates)
    duplicate_schedules = schedule_df[schedule_df.duplicated(subset=['SCH_SCHED_NBR'])]
    if not duplicate_schedules.empty:
        print("Duplicate schedules in ", site, ". Duplicate schedules are:")
        print(f"{duplicate_schedules.shape[0]} records for {duplicate_schedules}")
    # closing connection so minimum time connected
    conn.close()
    print('Connection closed')
    return schedule_df, stops_df, site, missing_schedule_data, types_dates


def validation(schedules_df, stops_df, site, zb, missing_schedule_data, TypesDates):

    print("Validating schedules...")
    zb_df = pd.read_excel(zb, sheet_name='Miles & Hours')

    original_schedules = schedules_df
    original_stops = stops_df

    # clean zero base df
    header_row = 13
    new_header = zb_df.iloc[header_row]
    zb_df = zb_df[header_row + 1:]
    zb_df.columns = new_header.values
    zb_df = zb_df.dropna(how='all')
    zb_df = zb_df.replace(r'^\s*$', np.nan, regex = True)
    zb_df = zb_df.fillna(0)
    zb_df = zb_df.iloc[:, 9:18]
    zb_df = zb_df[zb_df['Annual Work Hours'] != 0].reset_index(drop=True)
    zb_df['FREQ'] = zb_df['FREQ'].map(str)
    zb_df['FREQ'] = zb_df['FREQ'].apply(lambda x: x.zfill(4))
    zb_df['FREQ'] = zb_df['FREQ'].apply(lambda x: x.upper())
    zb_df['Schedule #'] = zb_df['Schedule #'].map(str)



    #set revelant zb info
    list_of_zb_sched = list(zb_df['Schedule #'])
    all_zb_sched_info = {}
    for sched in list_of_zb_sched:
        temp_df = zb_df[zb_df['Schedule #']== sched].reset_index(drop=True)
        type = list(temp_df['Type (D / DS / I / F)'])[0]
        freq = list(temp_df['FREQ'])[0]
        old_time = list(temp_df['Paid Dly Hrs'])[0]
        split_time = old_time.split(':')
        hours = int(split_time[0])
        minutes = int(split_time[1])/60
        new_time = hours+minutes
        vehicle = list(temp_df['M/T/S'])[0]
        if vehicle == 'S':
            vehicle = 'T'
        all_zb_sched_info[sched] = [freq, old_time, new_time, type, vehicle]

    #find lunch duration
    stops_df['STOP_NAME'] = stops_df['STOP_NAME'].apply(lambda x: x.lower())
    lunch_stops = stops_df[stops_df['STOP_NAME']=='lunch'].reset_index(drop = True)
    lunch_stops['DEP_TIME'] = lunch_stops['DEP_TIME'].astype('datetime64[ns]')
    lunch_stops['ARR_TIME'] = lunch_stops['ARR_TIME'].astype('datetime64[ns]')
    lunch_stops['Lunch Duration'] = lunch_stops['DEP_TIME']- lunch_stops['ARR_TIME']
    lunch_stops_final = lunch_stops[['SCH_SCHED_NBR', 'Lunch Duration']]

    #check for stops over 8 hours
    stops_df['DEP_TIME'] = stops_df['DEP_TIME'].astype('datetime64[ns]')
    stops_df['ARR_TIME'] = stops_df['ARR_TIME'].astype('datetime64[ns]')
    stops_df['Stop Duration'] = stops_df['DEP_TIME'] - stops_df['ARR_TIME']

    stops_no_lunch = stops_df[stops_df['STOP_NAME']!='lunch'].reset_index(drop=True)
    time_check1 = pd.to_timedelta('0 days 08:00:00')
    time_check2 = pd.to_timedelta('-1 days +08:00:00')
    time_check3 = pd.to_timedelta('0 days 00:00:00')
    over_8_stops_df1 = stops_no_lunch[stops_no_lunch['Stop Duration']>= time_check1].reset_index(drop=True)
    over_8_stops_df2 = stops_no_lunch.loc[(stops_no_lunch['Stop Duration'] > time_check2) & (stops_no_lunch['Stop Duration']<time_check3)].reset_index(drop=True)
    over_8_stops_df = pd.concat([over_8_stops_df1,over_8_stops_df2])
    over_8_stops_df.reset_index(drop=True,inplace=True)
    over_8_hr_stops = {}
    stops_over_8 = []

    if not over_8_stops_df.empty:
        for i in range(over_8_stops_df.shape[0]):
            sched_num = over_8_stops_df.loc[i,'SCH_SCHED_NBR']
            stop_num = over_8_stops_df.loc[i, 'STOP_NBR']
            stop_time = over_8_stops_df.loc[i, 'Stop Duration']
            stop_time_br = str(stop_time).split(' ')
            stop_time = stop_time_br[2]
            stop_time = stop_time[:6]
            if '+' in stop_time:
                stop_time = stop_time.strip('+')
            else:
                stop_time = stop_time.strip(':')
            over_8_hr_stops[sched_num] = [stop_num, stop_time]
            stops_over_8.append(sched_num)
    #check for stop time overlap
    problem_stops = pd.DataFrame()
    problem_stops = stop_time_mismatch(stops_df)

    #add lunch duration to schedules df
    i = 0
    for schedule in schedules_df['SCH_SCHED_NBR']:
        temp_stops = lunch_stops_final[lunch_stops_final['SCH_SCHED_NBR']== schedule]
        try:
            lunch_duration = list(temp_stops['Lunch Duration'])[0]
        except:
            lunch_duration  = pd.to_timedelta('0 days 00:00:00')
        schedules_df.loc[i, 'LUNCH TIME'] = lunch_duration
        i +=1

    schedules_df['SCH_DURATION'] = pd.to_datetime(schedules_df['SCH_DURATION'])
    schedules_df['Paid Hours'] = schedules_df['SCH_DURATION']-schedules_df['LUNCH TIME']
    schedules_df['Paid Hours'] = schedules_df['Paid Hours'].apply(lambda x: x.strftime('%H:%M'))
    schedules_df['SCH_SCHED_NBR'] = schedules_df['SCH_SCHED_NBR'].map(str)

    #set relevant VITAL info
    list_of_VITAL_schedules = list(schedules_df['SCH_SCHED_NBR'])
    all_VITAL_sched_info = {}
    for sched in list_of_VITAL_schedules:
        temp_df =schedules_df[schedules_df['SCH_SCHED_NBR']==sched].reset_index(drop=True)
        freq = list(temp_df['FRQ_CD'])[0]
        old_time = list(temp_df['Paid Hours'])[0]
        split_time = old_time.split(':')
        hours = int(split_time[0])
        minutes = int(split_time[1])/60
        new_time = hours + minutes
        tractor = list(temp_df['TRACTOR_IND'])[0]
        if tractor == 'Y':
            vehicle = 'T'
        elif tractor == 'N':
            vehicle = 'M'
        all_VITAL_sched_info[sched] = [freq, old_time, new_time, vehicle]

    #missing/extra schedules
    missing = list(set(list_of_zb_sched).difference(list_of_VITAL_schedules))
    extras = list(set(list_of_VITAL_schedules).difference(list_of_zb_sched))

    #schedules with no type in zb form
    no_type = zb_df[zb_df['Type (D / DS / I / F)']== 0 ]
    no_type_scheds = list(no_type['Schedule #'])

    #schedules with 0 mileage and not tractors - VITAL
    zero_mileage_df_v = schedules_df[schedules_df['MILEAGE_NBR'] ==0]
    zero_mileage_incorrect_vehicle_df_v = zero_mileage_df_v[zero_mileage_df_v['TRACTOR_IND'] != 'Y']
    fix_vehicle_schedules_v = list(zero_mileage_incorrect_vehicle_df_v['SCH_SCHED_NBR'])

    # schedules with 0 mileage and not tractors - ZB
    zero_mileage_df_zb = zb_df[zb_df['Dly Mileage'] == 0]
    zero_mileage_incorrect_vehicle_df_zb = zero_mileage_df_zb[zero_mileage_df_zb['M/T/S'] == 'M']
    fix_vehicle_schedules_zb = list(zero_mileage_incorrect_vehicle_df_zb['Schedule #'])

    #for i in missing:
     #   del all_zb_sched_info[i]

    for i in no_type_scheds:
        missing.remove(i)

    for i in extras:
        del all_VITAL_sched_info[i]

    #check freqs
    no_match_freqs = []
    for sched in all_zb_sched_info.keys():
        if sched in missing or sched in no_type_scheds:
            continue
        else:
            if all_zb_sched_info[sched][0] != all_VITAL_sched_info[sched][0]:
                no_match_freqs.append(sched)

    #check vehicles
    no_match_vehicles = []
    for sched in all_zb_sched_info.keys():
        if sched in missing or sched in no_type_scheds:
            continue
        else:
            if all_zb_sched_info[sched][4] != all_VITAL_sched_info[sched][3]:
                no_match_vehicles.append(sched)

    # check time
    no_match_time = []
    for sched in all_zb_sched_info.keys():
        if sched in missing or sched in no_type_scheds:
            continue
        else:
            if all_zb_sched_info[sched][2] != all_VITAL_sched_info[sched][2]:
                no_match_time.append(sched)

    #over 8 hour schedules
    over_8_zb = []
    for sched in all_zb_sched_info.keys():
        if sched in missing:
            continue
        else:
            if all_zb_sched_info[sched][2]>8:
                over_8_zb.append(sched)

    over_8_VITAL = []
    for sched in all_VITAL_sched_info.keys():
        if sched in missing:
            continue
        else:
            if all_VITAL_sched_info[sched][2]>8:
                over_8_VITAL.append(sched)

    #creating word document
    if len(no_match_freqs) > 0 or len(over_8_hr_stops)>0 or len(no_match_time) > 0 or len(over_8_zb)>0 or \
            len(over_8_VITAL)>0 or len(missing)>0 or len(extras)>0 or problem_stops.shape[0]>0 or len(no_type_scheds)>0 \
            or len(fix_vehicle_schedules_v)>0 or len(fix_vehicle_schedules_zb)>0 or len(no_match_vehicles)>0:
        doc = Document()
        header = doc.add_paragraph()
        header.add_run(site.upper()).bold = True
        if len(missing) >0:
            completely_missing = []
            different_type_date = []
            for i in missing:
                if missing_schedule_data.empty:
                    completely_missing.append(i)
                else:
                    missing_schedule_data['SCH_SCHED_NBR'] = missing_schedule_data['SCH_SCHED_NBR'].map(str)
                    missing_list = list(missing_schedule_data['SCH_SCHED_NBR'])
                    if i not in missing_list:
                        completely_missing.append(i)
                    else:
                        different_type_date.append(i)
            print(f"{len(missing)} schedules couldn't be pulled from VITAL")
            if len(completely_missing)>0:
                print(f"{len(completely_missing)} schedules don't exist in VITAL")
                doc.add_paragraph(f"The following {len(completely_missing)} schedules appear in the Zero Base Form but not in VITAL:")
                sched = doc.add_paragraph()
                completely_missing.sort()
                sched.add_run(completely_missing[0])
                for schedule in completely_missing[1:]:
                    sched.add_run(f", {schedule}")
                p1 = doc.add_paragraph()
                run = p1.add_run('ACTION: ')
                run.font.color.rgb = RGBColor(255,0,0)
                p1.add_run(f'Please add the {len(completely_missing)} schedules to VITAL.')
                doc.add_paragraph()
            if len(different_type_date)>0:
                print(f"{len(different_type_date)} schedules are in VITAL but have different type and/or effective date")
                doc.add_paragraph(f"The following {len(different_type_date)} schedules have different type and/or effective date in VITAL and the Zero Base Form: "
                                  f"(VITAL type - date, ZB type - date)")
                different_type_date.sort()
                for schedule in different_type_date:
                    temp_df = missing_schedule_data[missing_schedule_data['SCH_SCHED_NBR']==schedule]
                    zb_type = all_zb_sched_info[schedule][3]
                    zb_date = TypesDates[zb_type]
                    if zb_type == 'I':
                        zb_date = 'Now'
                    else:
                        zb_date = zb_date.replace("-","/")
                        zb_date_split = zb_date.split('/')
                        zb_date = zb_date_split[1] +'/' + zb_date_split[0] + '/' + zb_date_split[2]
                    records = temp_df.shape[0]
                    if records ==1:
                        v_type = list(temp_df['DECODE_DESC'])[0]
                        v_date = list(temp_df['SCH_EFFECT_DTM'])[0]
                        v_date = v_date.strftime("%m/%d/%Y")
                        doc.add_paragraph(f'{schedule} ({v_type} - {v_date}, {zb_type} - {zb_date})', style = 'List Bullet')
                    else:
                        doc.add_paragraph(f'{schedule} has {records} entries in VITAL. None of them match the zero base ({zb_type} - {zb_date})', style = 'List Bullet')
                        list_of_v_types = list(temp_df['DECODE_DESC'])
                        list_of_v_dates = list(temp_df['SCH_EFFECT_DTM'])
                        for i in range(len(list_of_v_dates)):
                            date = list_of_v_dates[i]
                            list_of_v_dates[i] = date.strftime("%m/%d/%Y")
                        for i in range(len(list_of_v_types)):
                            doc.add_paragraph(f'{list_of_v_types[i]} - {list_of_v_dates[i]}', style = 'List Bullet 2')
                p1 = doc.add_paragraph()
                run = p1.add_run('ACTION: ')
                run.font.color.rgb = RGBColor(255, 0, 0)
                p1.add_run(f'Please correct the {len(different_type_date)} schedules in VITAL or change them on the Zero Base form.')
                doc.add_paragraph()
                #     sched.add_run(completely_missing[0])
                # for schedule in completely_missing[1:]:
                #     sched.add_run(f", {schedule}")
                # p1 = doc.add_paragraph()
                # run = p1.add_run('ACTION: ')
                # run.font.color.rgb = RGBColor(255,0,0)
                # p1.add_run(f'Please add the {len(missing)} schedules to VITAL.')
                # doc.add_paragraph
        if len(no_type_scheds)>0:
            print(f"{len(no_type_scheds)} schedules don't have a type listed in the Zero Base Form.")
            doc.add_paragraph(f'The following {len(no_type_scheds)} schedules appear in the Zero Base Form without a type but list mileage/workhours:')
            sched = doc.add_paragraph()
            no_type_scheds.sort()
            sched.add_run(no_type_scheds[0])
            for schedule in no_type_scheds[1:]:
                sched.add_run(f", {schedule}")
            p1 = doc.add_paragraph()
            run = p1.add_run('ACTION: ')
            run.font.color.rgb = RGBColor(255, 0, 0)
            p1.add_run(f'Please input the type of each schedule in the Zero Base Form (I/D/DS/F).')
            doc.add_paragraph()
        if len(extras) >0:
            print(f'{len(extras)} extra schedules in VITAL')
            doc.add_paragraph(f'The following {len(extras)} schedules appear in VITAL but not in the Zero Base Form:')
            sched =doc.add_paragraph()
            extras.sort()
            sched.add_run(extras[0])
            for schedule in extras[1:]:
                sched.add_run(f", {schedule}")
            p1 = doc.add_paragraph()
            run = p1.add_run('ACTION: ')
            run.font.color.rgb = RGBColor(255,0,0)
            p1.add_run(f'Please confirm the {len(extras)} schedules do not get used in the Zero Base evaulation.')
            doc.add_paragraph()
        if len(no_match_freqs)>0:
            print(f'{len(no_match_freqs)} schedules have mismatch frequencies')
            doc.add_paragraph(f'The following {len(no_match_freqs)} schedules have different frequencies in VITAL and '
                              f'on the Zero Base Form: (VITAL FREQ, ZB FREQ)')
            nm_f = doc.add_paragraph()
            no_match_freqs.sort()
            sch = no_match_freqs[0]
            nm_f.add_run(f'{sch} ({all_VITAL_sched_info[sch][0]}, {all_zb_sched_info[sch][0]})')
            for schedule in no_match_freqs[1:]:
                nm_f.add_run(f', {schedule} ({all_VITAL_sched_info[schedule][0]}, {all_zb_sched_info[schedule][0]})')
            p1 = doc.add_paragraph()
            run = p1.add_run('ACTION: ')
            run.font.color.rgb = RGBColor(255, 0, 0)
            p1.add_run('Please edit either the VITAL schedules or the post zero base form, so that the frequencies match.')
            doc.add_paragraph()
        if len(no_match_time)>0:
            print(f'{len(no_match_time)} schedules have mismatch paid time')
            doc.add_paragraph(f'The following {len(no_match_time)} schedules have different paid time in VITAL and '
                              f'on the Zero Base Form: (VITAL Paid Time, ZB Paid Time)')
            nm_t = doc.add_paragraph()
            no_match_time.sort()
            sch = no_match_time[0]
            nm_t.add_run(f'{sch} ({all_VITAL_sched_info[sch][1]}, {all_zb_sched_info[sch][1]})')
            for schedule in no_match_time[1:]:
                nm_t.add_run(f', {schedule} ({all_VITAL_sched_info[schedule][1]}, {all_zb_sched_info[schedule][1]})')
            p1 = doc.add_paragraph()
            run = p1.add_run('ACTION: ')
            run.font.color.rgb = RGBColor(255, 0, 0)
            p1.add_run('Please edit either the VITAL schedules or the post zero base form, so that the paid time match.')
            doc.add_paragraph()
        if len(no_match_vehicles)>0:
            print(f'{len(no_match_vehicles)} schedules have mismatch vehicle type')
            doc.add_paragraph(f'The following {len(no_match_vehicles)} schedules have different vehicle types in VITAL and '
                              f'on the Zero Base Form: (VITAL Vehicle, ZB Vehicle)')
            nm_v = doc.add_paragraph()
            no_match_vehicles.sort()
            sch = no_match_vehicles[0]
            nm_v.add_run(f'{sch} ({all_VITAL_sched_info[sch][3]}, {all_zb_sched_info[sch][4]})')
            for schedule in no_match_vehicles[1:]:
                nm_v.add_run(f', {schedule} ({all_VITAL_sched_info[schedule][3]}, {all_zb_sched_info[schedule][4]})')
            p1 = doc.add_paragraph()
            run = p1.add_run('ACTION: ')
            run.font.color.rgb = RGBColor(255, 0, 0)
            p1.add_run('Please edit either the VITAL schedules or the post zero base form, so that the vehicle type matches.'
                       '(Spotter schedules should be marked as Tractors in VITAL)')
            doc.add_paragraph()
        if len(over_8_zb)>0:
            print(f'{len(over_8_zb)} schedules have paid time greater than 8 hours in ZB')
            doc.add_paragraph(f'The following {len(over_8_zb)} schedules have over 8 hours paid time in the Zero Base Form: (Paid Time)')
            o8z = doc.add_paragraph()
            over_8_zb.sort()
            o8z.add_run(f'{over_8_zb[0]} ({all_zb_sched_info[over_8_zb[0]][1]})')
            for schedule in over_8_zb[1:]:
                o8z.add_run(f', {schedule} ({all_zb_sched_info[schedule][1]})')
            p1 = doc.add_paragraph()
            run = p1.add_run('ACTION: ')
            run.font.color.rgb = RGBColor(255, 0, 0)
            p1.add_run(f'Please correct the {len(over_8_zb)} schedules paid time in the Zero Base Form to be 8 hours or less.')
            doc.add_paragraph()
        if len(over_8_VITAL)>0:
            print(f'{len(over_8_VITAL)} schedules have paid time greater than 8 hours in VITAL')
            doc.add_paragraph(f'The following {len(over_8_VITAL)} schedules have over 8 hours paid time in VITAL: (Paid Time)')
            o8v = doc.add_paragraph()
            over_8_VITAL.sort()
            o8v.add_run(f'{over_8_VITAL[0]} ({all_VITAL_sched_info[over_8_VITAL[0]][1]})')
            for schedule in over_8_VITAL[1:]:
                o8v.add_run(f', {schedule} ({all_VITAL_sched_info[schedule][1]})')
            p1 = doc.add_paragraph()
            run = p1.add_run('ACTION: ')
            run.font.color.rgb = RGBColor(255, 0, 0)
            p1.add_run(f'Please correct the {len(over_8_VITAL)} schedules paid time in VITAL to be 8 hours or less.')
            doc.add_paragraph()
        if len(stops_over_8)>0:
            print(f'{len(stops_over_8)} schedules have a stop greater than 8 hours in VITAL')
            doc.add_paragraph(
                f'The following {len(stops_over_8)} schedules have a stop greater than 8 hours in VITAL: (Stop Number, Stop Duration H:M)')
            o8s = doc.add_paragraph()
            stops_over_8.sort()
            o8s.add_run(f'{stops_over_8[0]} ({over_8_hr_stops[stops_over_8[0]][0]}, {over_8_hr_stops[stops_over_8[0]][1]})')
            for schedule in stops_over_8[1:]:
                o8s.add_run(f', {schedule} ({over_8_hr_stops[schedule][0]}, {over_8_hr_stops[schedule][1]})')
            p1 = doc.add_paragraph()
            run = p1.add_run('ACTION: ')
            run.font.color.rgb = RGBColor(255, 0, 0)
            p1.add_run(f'Please correct the {len(stops_over_8)} schedules in VITAL.')
            doc.add_paragraph()
        if len(fix_vehicle_schedules_v)>0:
            print(f"{len(fix_vehicle_schedules_v)} schedules have the wrong vehicle type listed in VITAL.")
            doc.add_paragraph(f"The following {len(fix_vehicle_schedules_v)} schedules have a listed annual mileage of 0 but don't have the tractor indictor marked in VITAL:")
            sched = doc.add_paragraph()
            fix_vehicle_schedules_v.sort()
            sched.add_run(fix_vehicle_schedules_v[0])
            for schedule in fix_vehicle_schedules_v[1:]:
                sched.add_run(f", {schedule}")
            p1 = doc.add_paragraph()
            run = p1.add_run('ACTION: ')
            run.font.color.rgb = RGBColor(255, 0, 0)
            p1.add_run(f'Please mark the tractor indicator in VITAL or change the mileage for the above {len(fix_vehicle_schedules_v)} schedules.')
            doc.add_paragraph()
        if len(fix_vehicle_schedules_zb)>0:
            print(f"{len(fix_vehicle_schedules_zb)} schedules have the wrong vehicle type listed in the ZB Form.")
            doc.add_paragraph(f"The following {len(fix_vehicle_schedules_zb)} schedules have a listed annual mileage of 0 but are marked as vehicle type 'M' in the Zero Base form:")
            sched = doc.add_paragraph()
            fix_vehicle_schedules_zb.sort()
            sched.add_run(fix_vehicle_schedules_zb[0])
            for schedule in fix_vehicle_schedules_zb[1:]:
                sched.add_run(f", {schedule}")
            p1 = doc.add_paragraph()
            run = p1.add_run('ACTION: ')
            run.font.color.rgb = RGBColor(255, 0, 0)
            p1.add_run(f'Please change the vehicle to T or S in the Zero Base Form or change the mileage for the above {len(fix_vehicle_schedules_zb)} schedules.')
            doc.add_paragraph()
        if problem_stops.shape[0]>0:
            print(f"Something is wrong with {int(problem_stops.shape[0]/2)} schedules in VITAL")
            doc.add_paragraph(f"The following {int(problem_stops.shape[0]/2)} schedules have stops overlapping in time")
            list_of_schedules = set(list(problem_stops['Schedule #']))
            for stop_sched in list_of_schedules:
                temp_df = problem_stops[problem_stops['Schedule #']==stop_sched]
                occurences = set(list(temp_df['Indicator']))
                num_occur = len(occurences)
                for i in occurences:
                    schedules_occurence = temp_df[temp_df['Indicator']==i].reset_index(drop=True)
                    first_stop = schedules_occurence.loc[0, 'Stop #']
                    first_start = schedules_occurence.loc[0, 'Start Time']
                    first_end = schedules_occurence.loc[0, 'End Time']
                    second_stop = schedules_occurence.loc[1, 'Stop #']
                    second_start = schedules_occurence.loc[1, 'Start Time']
                    second_end = schedules_occurence.loc[1, 'End Time']
                    doc.add_paragraph(f"Schedule {stop_sched} - Stop {first_stop} ({first_start}-{first_end}) overlaps with Stop {second_stop} ({second_start}-{second_end})", style='List Bullet')


        doc.save(str('DataFiles/' + site.upper() + ' Questions ' + time.strftime("%m.%d.%y") +'.docx'))
        print('Send questions')
        #os.remove('DataFiles\\' + site + ' SCHEDULES ZB ' + time.strftime("%Y%m%d") + '.xlsx')
        #os.remove('DataFiles\\' + site + ' STOPS ZB ' + time.strftime("%Y%m%d") + '.xlsx')
    else:
        print("Everything looks good. Schedules and Stops printed to Datafiles")
        original_schedules['SCH_DURATION'] = original_schedules['SCH_DURATION'].apply(lambda  x: x.strftime('%H:%M'))
        original_schedules = original_schedules[['AREA_NAME', 'PVS_SITE_ID', 'SITE_NAME', 'DECODE_DESC', 'SCH_SCHED_NBR',
            'TOUR_NBR', 'SCHED_TYPE_ID', 'RUN_NBR', 'FRQ_CD', 'TRACTOR_IND',
            'VEH_1_ID', 'VEH_2_ID', 'VEH_3_ID', 'MILEAGE_NBR', 'START_TIME',
            'END_TIME', 'SCH_DURATION', 'TOT_STOP_CNT', 'SCH_EFFECT_DTM', 'END_DT',
            'SCH_ROUTE_ID', 'TRIP_CD', 'ROUTE_PREV_IND', 'SCHED_NO_PREV_IND',
            'END_TIMES_DT', 'INV_TIMES_IND', 'TIME_TRIP_ID', 'UNSCHEDULE_TRIP_IND',
            'SHOW_START_TIME_IND', 'COMMENT_TEXT', 'VER_NBR']]
        #original_schedules.to_excel('DataFiles\\' + site + ' SCHEDULES ZB ' + time.strftime("%Y%m%d") + '.xlsx', index=False)
        #original_stops.to_excel('DataFiles\\' + site + ' STOPS ZB ' + time.strftime("%Y%m%d") + '.xlsx', index=False)


# gets all info on missing schedules that is currently listed in VITAL
def process_missing_schedules(site, not_found_schedules, conn):
    check_vital_query = ScheduleQuery()
    missing_sched_query_str = check_vital_query.build_missing_sched_query(site)
    # pulls all schedule information from VITAL for one site
    all_schedules = pd.read_sql(missing_sched_query_str, conn)
    missing_schedules = []
    schedules_not_in_VITAL = []
    for schedule in not_found_schedules:
        # checks if the schedule number that is missing is in the df of all schedules for a site
        schedule_in_vital = all_schedules[(all_schedules['SCH_SCHED_NBR'] == str(schedule))]
        if schedule_in_vital.empty:
            schedules_not_in_VITAL.append(str(schedule))
            print('Schedule ', schedule, 'is not in VITAL!')
        else:
            # if info on schedule number is found, add it to the list to print with missing schedules
            missing_schedules.append(schedule_in_vital)
    # combining all dfs for missing schedules found
    final_missing_schedules = pd.DataFrame()
    if len(missing_schedules)!=0:
        final_missing_schedules = pd.concat(missing_schedules)
        # printing all missing schedules
        final_missing_schedules.to_excel(str('DataFiles/' + site + '_MissingSchedules.xlsx'), index=False)
    if len(schedules_not_in_VITAL)!=0:
        not_in_VITAL_df = pd.DataFrame(schedules_not_in_VITAL, columns=['Schedules not in VITAL'])
        not_in_VITAL_df.to_excel(str('DataFiles/' + site + '_SchedulesNotinVITAL.xlsx'), index = False)
    # printing all schedules for 1 site (for now, for reference)
    all_schedules.to_excel(str('DataFiles/' + site + '_AllSchedules.xlsx'), index=False)
    return final_missing_schedules


if __name__ == "__main__":
    #sort_and_print_sched_stops('Zero Base Forms/2021 LOS ANGELES NDC Zero Base Forms v062121 003 71.2021.xlsm')
    sort_and_print_sched_stops('Zero Base Forms/2021 ATLANTA Zero Base Forms v061521 2.xlsm')